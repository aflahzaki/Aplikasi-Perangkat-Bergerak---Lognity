from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def ss(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'E0E0E0')
    run._element.get_or_add_rPr().append(shd)

def center(text, size=11, bold=False, color=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor(*color)
    return p

# COVER
for _ in range(3):
    doc.add_paragraph()
center('[LOGO UNIVERSITAS TELKOM]', 14, True)
doc.add_paragraph()
center('UNIVERSITAS TELKOM', 16, True)
center('FAKULTAS TEKNIK DAN INFORMATIKA', 14, True)
for _ in range(2):
    doc.add_paragraph()
center('LAPORAN UJIAN AKHIR SEMESTER', 13, True)
center('MATA KULIAH APLIKASI PERANGKAT BERGERAK', 13, True)
for _ in range(2):
    doc.add_paragraph()
center('LOGNITY - Platform E-Library & Forum Diskusi Kampus', 15, True, (0x1A, 0x23, 0x7E))
for _ in range(3):
    doc.add_paragraph()
center('Disusun Oleh:', 12, True)
members = [
    ('Aresky Brilliant Sombolinggi', '1301210001'),
    ('Aflah Zaki Siregar', '1301210002'),
    ('Stiefanny Dwi Chandra', '1301210003'),
    ('Elsa Melisa Silaen', '1301210004'),
]
for name, nim in members:
    center(f'{name} ({nim})', 11)
doc.add_page_break()

# BAB 1 - PENDAHULUAN
doc.add_heading('BAB 1 - PENDAHULUAN', level=1)

doc.add_heading('1.1 Latar Belakang', level=2)
doc.add_paragraph(
    'Di era digital saat ini, kebutuhan mahasiswa akan platform kolaborasi dan berbagi materi perkuliahan '
    'semakin meningkat. Mahasiswa memerlukan wadah yang memudahkan mereka untuk saling berbagi catatan, '
    'berdiskusi tentang materi kuliah, dan mengakses referensi pembelajaran secara mudah dan terstruktur. '
    'Namun, platform yang ada saat ini masih terpisah-pisah dan belum terintegrasi dengan baik.'
)
doc.add_paragraph(
    'Lognity hadir sebagai solusi berupa platform mobile yang mengintegrasikan fitur Forum Diskusi '
    'dan E-Library dalam satu aplikasi. Dengan pendekatan gamifikasi (XP, Level, dan Leaderboard), '
    'Lognity mendorong partisipasi aktif mahasiswa dalam berbagi pengetahuan dan berkolaborasi '
    'secara akademik.'
)

doc.add_heading('1.2 Rumusan Masalah', level=2)
doc.add_paragraph('Berdasarkan latar belakang di atas, rumusan masalah yang diangkat adalah:')
doc.add_paragraph('1. Keterbatasan akses materi perkuliahan yang tersebar di berbagai platform berbeda.', style='List Number')
doc.add_paragraph('2. Kurangnya platform diskusi akademik yang terstruktur antar mahasiswa.', style='List Number')
doc.add_paragraph('3. Rendahnya motivasi mahasiswa untuk berbagi materi karena tidak adanya sistem reward.', style='List Number')

doc.add_heading('1.3 Tujuan', level=2)
doc.add_paragraph('Tujuan dari pengembangan aplikasi Lognity adalah:')
doc.add_paragraph('1. Membangun platform mobile yang mengintegrasikan Forum Diskusi dan E-Library.', style='List Number')
doc.add_paragraph('2. Menyediakan sistem gamifikasi (XP, Level, Leaderboard) untuk mendorong partisipasi aktif.', style='List Number')
doc.add_paragraph('3. Memudahkan mahasiswa dalam mengakses, berbagi, dan mendiskusikan materi perkuliahan.', style='List Number')

doc.add_page_break()

# BAB 2 - PEMBAGIAN TUGAS
doc.add_heading('BAB 2 - PEMBAGIAN TUGAS DI TIM', level=1)
doc.add_paragraph('Akan dilengkapi kemudian.')
doc.add_page_break()

# BAB 3 - ARSITEKTUR APLIKASI
doc.add_heading('BAB 3 - ARSITEKTUR APLIKASI', level=1)

doc.add_heading('3.1 Tech Stack', level=2)
doc.add_paragraph('Aplikasi Lognity dibangun menggunakan tech stack berikut:')
doc.add_paragraph('Mobile: Flutter (Dart) SDK ^3.11.0', style='List Bullet')
doc.add_paragraph('Backend: Laravel (PHP) dengan database SQLite/MySQL', style='List Bullet')
doc.add_paragraph('API: RESTful API (Laravel Sanctum token-based authentication)', style='List Bullet')

doc.add_heading('3.2 Package yang Digunakan (pubspec.yaml)', level=2)
doc.add_paragraph('Berikut adalah daftar package yang digunakan dalam pengembangan aplikasi mobile:')

table = doc.add_table(rows=9, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Package'
hdr[1].text = 'Versi'
hdr[2].text = 'Fungsi'
for cell in hdr:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

packages = [
    ('cupertino_icons', '1.0.8', 'Ikon gaya iOS untuk konsistensi UI'),
    ('http', '1.6.0', 'HTTP client untuk komunikasi dengan REST API backend'),
    ('provider', '6.1.5+1', 'State Management menggunakan ChangeNotifier pattern'),
    ('shared_preferences', '2.5.5', 'Penyimpanan data lokal (token auth, preferensi tema)'),
    ('url_launcher', '6.3.2', 'Membuka URL/file di browser eksternal'),
    ('file_picker', '11.0.2', 'Pemilihan file untuk upload lampiran forum'),
    ('google_fonts', '6.2.1', 'Font Poppins untuk konsistensi typography UI'),
    ('flutter_svg', '2.0.10', 'Render SVG assets (logo dan ikon)'),
]
for i, (pkg, versi, fungsi) in enumerate(packages):
    row = table.rows[i + 1].cells
    row[0].text = pkg
    row[1].text = versi
    row[2].text = fungsi

doc.add_heading('3.3 Latar Belakang Pemilihan Database', level=2)
doc.add_paragraph(
    'SQLite digunakan untuk environment development karena ringan dan tidak memerlukan konfigurasi '
    'server database terpisah, sehingga mempercepat proses pengembangan. MySQL digunakan untuk '
    'production karena mendukung skalabilitas, concurrent access, dan fitur relational database '
    'yang lebih lengkap.'
)

doc.add_heading('3.4 Latar Belakang Pemilihan State Management', level=2)
doc.add_paragraph('Provider dipilih sebagai state management dengan pertimbangan:')
doc.add_paragraph('Terintegrasi baik dengan Flutter ecosystem dan direkomendasikan oleh tim Flutter.', style='List Bullet')
doc.add_paragraph('Implementasi sederhana dengan ChangeNotifier pattern.', style='List Bullet')
doc.add_paragraph('Cocok untuk aplikasi skala menengah seperti Lognity.', style='List Bullet')
doc.add_paragraph('Digunakan untuk ThemeProvider (dark/light mode toggle) dan manajemen state lainnya.', style='List Bullet')

doc.add_heading('3.5 Arsitektur Kode', level=2)
doc.add_paragraph('Struktur folder aplikasi mengikuti pola berikut:')
doc.add_paragraph('lib/ - Root folder source code', style='List Bullet')
doc.add_paragraph('  main.dart - Entry point aplikasi', style='List Bullet')
doc.add_paragraph('  screens/ - Halaman-halaman UI (login, register, dashboard, forum, dll)', style='List Bullet')
doc.add_paragraph('  services/ - Layer service (ApiService sebagai single point of contact ke backend)', style='List Bullet')
doc.add_paragraph('  theme/ - Konfigurasi tema (ThemeProvider untuk dark/light mode)', style='List Bullet')
doc.add_paragraph('  widgets/ - Komponen UI reusable', style='List Bullet')
doc.add_paragraph()
doc.add_paragraph('Pattern yang digunakan: Service-based architecture dimana ApiService bertindak sebagai single point of contact untuk semua komunikasi dengan backend REST API.')
doc.add_paragraph()
doc.add_paragraph('Routing menggunakan named routes:')
doc.add_paragraph('/onboarding - Halaman pengenalan aplikasi', style='List Bullet')
doc.add_paragraph('/login - Halaman masuk', style='List Bullet')
doc.add_paragraph('/register - Halaman pendaftaran', style='List Bullet')
doc.add_paragraph('/main - Halaman utama (Dashboard + Bottom Navigation)', style='List Bullet')

doc.add_page_break()

# BAB 4 - OUTPUT DAN FITUR APLIKASI
doc.add_heading('BAB 4 - OUTPUT DAN FITUR APLIKASI', level=1)
doc.add_paragraph('Bab ini menjelaskan secara detail setiap fitur yang tersedia dalam aplikasi Lognity beserta panduan screenshot untuk dokumentasi.')

# 4.1 Onboarding
doc.add_heading('4.1 Onboarding Screen', level=2)
doc.add_paragraph(
    'Aplikasi Lognity memiliki 3 halaman onboarding yang muncul saat pengguna pertama kali '
    'membuka aplikasi. Onboarding bertujuan memperkenalkan fitur utama aplikasi.'
)
doc.add_heading('Halaman 1: "Belajar Lebih Seru"', level=3)
doc.add_paragraph('Menampilkan ikon rocket berwarna ungu (purple) dengan penjelasan bahwa Lognity adalah platform kolaborasi untuk mahasiswa. Pesan utama mengajak pengguna untuk belajar dengan cara yang lebih interaktif dan menyenangkan.')
doc.add_heading('Halaman 2: "Diskusi Terbuka"', level=3)
doc.add_paragraph('Menampilkan ikon forum berwarna pink dengan penjelasan fitur forum tanya jawab. Mahasiswa dapat bertanya, menjawab, dan berdiskusi secara terbuka.')
doc.add_heading('Halaman 3: "Kumpulkan XP & Level"', level=3)
doc.add_paragraph('Menampilkan ikon medal berwarna orange dengan penjelasan sistem gamifikasi. Setiap kontribusi memberikan XP yang meningkatkan level pengguna.')
doc.add_paragraph()
doc.add_paragraph('Navigasi: Dot indicator menunjukkan posisi halaman aktif. Tombol "Lanjut" untuk ke halaman berikutnya, berubah menjadi "Mulai" di halaman terakhir. Tombol "Lewati" tersedia untuk skip langsung ke login.')

ss('[Screenshot 4.1.1: Onboarding Halaman 1 - Buka aplikasi pertama kali (atau clear data aplikasi). Tampilan pertama yang muncul adalah Onboarding dengan teks "Belajar Lebih Seru" dan ikon roket berwarna ungu. Ambil screenshot seluruh layar.]')
ss('[Screenshot 4.1.2: Onboarding Halaman 2 - Swipe ke kanan atau tap tombol "Lanjut". Tampilan "Diskusi Terbuka" dengan ikon forum berwarna pink.]')
ss('[Screenshot 4.1.3: Onboarding Halaman 3 - Swipe ke kanan lagi. Tampilan "Kumpulkan XP & Level" dengan ikon medal berwarna orange dan tombol berubah menjadi "Mulai".]')

# 4.2 Login
doc.add_heading('4.2 Login Screen', level=2)
doc.add_paragraph(
    'Halaman login menampilkan background gradient dari biru ke ungu dengan dekoratif circles. '
    'Di bagian tengah atas terdapat logo Lognity berupa huruf "L" dengan gradient biru-ungu. '
    'Form login terdiri dari field Email dan Password (dengan toggle visibility untuk melihat/sembunyikan password). '
    'Tombol "Masuk" berwarna kuning mencolok sebagai CTA utama. '
    'Di bawah tombol terdapat link "Belum punya akun? Daftar sekarang" untuk navigasi ke halaman registrasi.'
)
ss('[Screenshot 4.2: Login Screen - Setelah menyelesaikan onboarding atau tap "Lewati", Anda akan diarahkan ke halaman Login. Tampilan form login dengan background gradient biru-ungu, logo "L" di atas, field Email dan Password, serta tombol kuning "Masuk".]')

# 4.3 Register
doc.add_heading('4.3 Register Screen', level=2)
doc.add_paragraph(
    'Halaman registrasi memiliki desain serupa dengan login screen. '
    'Form pendaftaran terdiri dari 4 field: Username, Email, Password, dan Konfirmasi Password. '
    'Tombol "Daftar" berwarna kuning sebagai CTA. '
    'Di bawah terdapat link "Sudah punya akun? Masuk di sini" untuk kembali ke login.'
)
ss('[Screenshot 4.3: Register Screen - Dari halaman Login, tap link "Daftar sekarang" di bagian bawah. Halaman pendaftaran akan muncul dengan 4 field: Username, Email, Password, dan Konfirmasi Password, serta tombol "Daftar".]')

# 4.4 Dashboard
doc.add_heading('4.4 Dashboard (Main Screen - Tab Tengah)', level=2)
doc.add_paragraph('Dashboard adalah halaman utama setelah login berhasil. Terdiri dari beberapa komponen:')
doc.add_heading('AppBar', level=3)
doc.add_paragraph('Logo Lognity + nama "Lognity" di kiri, toggle dark mode (ikon bulan/matahari), ikon chat, dan ikon profil di kanan.')
doc.add_heading('Hero Banner', level=3)
doc.add_paragraph('Card gradient biru-ungu menampilkan foto profil pengguna, sapaan "Halo, [username]!", serta pills untuk role dan level pengguna.')
doc.add_heading('XP Card', level=3)
doc.add_paragraph('Menampilkan total poin XP dengan progress bar visual menuju level berikutnya.')
doc.add_heading('Status Card', level=3)
doc.add_paragraph('Menampilkan role aktif pengguna dan status keanggotaan (Starter/Member/Elite).')
doc.add_heading('Kuota Harian', level=3)
doc.add_paragraph('Menampilkan sisa kuota harian: Request Forum dan Interaksi & Komentar dengan progress bar.')
doc.add_heading('Quick Access Cards', level=3)
doc.add_paragraph('Card Forum Diskusi (gradient pink ke merah) dengan jumlah diskusi aktif. Card E-Library (gradient hijau). Card Papan Peringkat (gradient orange) dengan emoji trophy. Card "Topik Sedang Hangat" (api emoji).')
doc.add_heading('Bottom Navigation', level=3)
doc.add_paragraph('3 tab: Forum (kiri) | Dashboard (tengah) | E-Library (kanan).')

ss('[Screenshot 4.4.1: Dashboard - Bagian Atas - Setelah login berhasil, Anda langsung masuk ke Dashboard (tab tengah di bottom navigation). Screenshot bagian atas: Hero banner gradient biru-ungu dengan foto profil, sapaan nama, role & level badges.]')
ss('[Screenshot 4.4.2: Dashboard - Bagian Tengah - Scroll ke bawah sedikit. Terlihat card XP (Total Poin + progress bar), card Status Akun, dan Kuota Harian.]')
ss('[Screenshot 4.4.3: Dashboard - Bagian Bawah - Scroll lebih ke bawah. Terlihat card Forum Diskusi (pink), E-Library (hijau), Papan Peringkat (orange), dan Topik Sedang Hangat.]')

# 4.5 Forum Diskusi
doc.add_heading('4.5 Forum Diskusi (Tab Kiri)', level=2)
doc.add_paragraph(
    'Halaman Forum Diskusi menampilkan daftar diskusi dari seluruh pengguna. '
    'Di bagian atas terdapat toggle Global/Mengikuti untuk filter diskusi. '
    'Search bar memungkinkan pencarian berdasarkan keyword. '
    'Ikon filter (tune) membuka panel filter dengan dropdown Kategori, Fakultas, dan Urutkan. '
    'Setiap card diskusi menampilkan: avatar pengguna, username, badge kategori, deskripsi singkat, '
    'jumlah upvote, dan jumlah komentar. Color accent bar di sisi kiri setiap card berwarna '
    'berdasarkan kategori diskusi. FAB "Diskusi Baru" tersedia di kanan bawah.'
)
ss('[Screenshot 4.5.1: Forum - Tampilan Utama - Tap tab "Forum" di bottom navigation (paling kiri). Tampilan daftar diskusi dengan toggle Global/Mengikuti di atas, search bar, dan list card diskusi berwarna-warni.]')
ss('[Screenshot 4.5.2: Forum - Filter Expanded - Tap ikon filter (tune) di samping search bar. Panel filter terbuka menampilkan dropdown Kategori, Fakultas, dan Urutkan.]')

# 4.6 Buat Diskusi Baru
doc.add_heading('4.6 Buat Diskusi Baru (Create Forum)', level=2)
doc.add_paragraph(
    'Form pembuatan diskusi baru terdiri dari: '
    'textarea untuk deskripsi request, dropdown Kategori, dropdown Fakultas, input Mata Kuliah, '
    'dan area upload file yang mendukung format PDF, JPG, PNG, ZIP, dan DOC. '
    'Tombol "Kirim Request" untuk mengirim diskusi baru.'
)
ss('[Screenshot 4.6: Buat Diskusi Baru - Dari halaman Forum, tap tombol FAB "Diskusi Baru" (tombol gradient biru-ungu di kanan bawah). Form pembuatan diskusi baru muncul dengan field deskripsi, dropdown kategori & fakultas, input mata kuliah, dan area upload file.]')

# 4.7 Detail Forum & Komentar
doc.add_heading('4.7 Detail Forum & Komentar', level=2)
doc.add_paragraph(
    'Halaman detail forum menampilkan pertanyaan lengkap beserta komentar/jawaban. '
    'AppBar memiliki judul "Detail Forum" dengan tombol edit (jika owner), delete, dan report. '
    'Card pertanyaan menampilkan avatar, username, badge kategori, deskripsi lengkap, lampiran (jika ada), '
    'dan tombol Upvote. Section "Komentar" menampilkan jumlah komentar dan daftar jawaban. '
    'Setiap komentar memiliki avatar, username, badge "Terbaik" (jika accepted), konten, '
    'lampiran, dan aksi (edit/delete/report). Owner forum dapat menerima jawaban dengan tombol '
    '"Terima Jawaban" (+50 Poin). Input bar di bawah layar berisi tombol attachment, text field, dan tombol send.'
)
ss('[Screenshot 4.7.1: Detail Forum - Tap salah satu card diskusi di halaman Forum. Halaman detail menampilkan pertanyaan lengkap di card atas dengan avatar, username, kategori badge, deskripsi, dan tombol Upvote.]')
ss('[Screenshot 4.7.2: Detail Forum - Komentar - Scroll ke bawah pada halaman detail. Section "Komentar" menampilkan jawaban-jawaban dari user lain. Jika ada jawaban yang diterima, akan ada badge hijau "Terbaik".]')
ss('[Screenshot 4.7.3: Detail Forum - Input Komentar - Lihat bagian paling bawah layar. Terdapat input bar dengan ikon attachment (penjepit kertas), text field "Tulis komentar...", dan tombol send gradient biru-ungu.]')

# 4.8 E-Library
doc.add_heading('4.8 E-Library (Tab Kanan)', level=2)
doc.add_paragraph(
    'E-Library menampilkan koleksi e-book dan materi pembelajaran. '
    'Search bar di atas untuk mencari berdasarkan judul atau penulis. '
    'Category chips (Semua, Buku, dst) untuk filter berdasarkan kategori. '
    'Toolbar menampilkan jumlah e-book dan toggle Grid/List view. '
    'Grid view: card dengan cover image, badge kategori, judul, dan nama penulis. '
    'List view: thumbnail di kiri dengan info buku di kanan. '
    'Tap pada item akan membuka file (PDF/dokumen) di browser eksternal.'
)
ss('[Screenshot 4.8.1: E-Library - Grid View - Tap tab "E-Library" di bottom navigation (paling kanan). Tampilan grid card buku dengan cover image, badge kategori, judul, dan nama penulis.]')
ss('[Screenshot 4.8.2: E-Library - List View - Tap ikon toggle view (grid/list) di toolbar kanan atas content. Tampilan berubah menjadi list horizontal dengan thumbnail cover di kiri dan info buku di kanan.]')

# 4.9 Profil Saya
doc.add_heading('4.9 Profil Saya', level=2)
doc.add_paragraph(
    'Halaman profil menampilkan informasi lengkap pengguna. '
    'Header gradient dengan foto profil besar di tengah, username, dan email. '
    'Stats menampilkan Level dan Poin. Badges horizontal list menampilkan pencapaian. '
    'Tab Postingan menampilkan list card postingan user di forum. '
    'Tab Interaksi menampilkan list card komentar yang pernah diberikan. '
    'Section Pengaturan berisi: Edit Profil, Ganti Password, dan Keluar (Logout).'
)
ss('[Screenshot 4.9.1: Profil - Bagian Atas - Dari Dashboard, tap ikon person (orang) di kanan atas AppBar. Halaman profil menampilkan header gradient dengan foto profil besar, username, email, dan stats Level & Poin.]')
ss('[Screenshot 4.9.2: Profil - Bagian Bawah - Scroll ke bawah. Terlihat section Badges, Postingan, Interaksi, dan menu Pengaturan (Edit Profil, Ganti Password, Keluar).]')
ss('[Screenshot 4.9.3: Edit Profil - Tap "Edit Profil" di section Pengaturan. Dialog muncul dengan field Username, Email, dan tombol "Pilih Foto" untuk ganti foto profil.]')

# 4.10 Direct Message
doc.add_heading('4.10 Direct Message (Chat)', level=2)
doc.add_paragraph(
    'Fitur chat memungkinkan komunikasi langsung antar pengguna. '
    'Daftar chat menampilkan avatar, username, pesan terakhir, waktu, dan badge unread. '
    'Swipe pada item chat untuk menghapus. '
    'Detail chat menggunakan bubble message (biru untuk pesan terkirim, putih/abu untuk diterima) dengan timestamp. '
    'Input bar di bawah memiliki tombol add, text field, dan tombol send berupa lingkaran gradient. '
    'Polling dilakukan setiap 3 detik untuk mengecek pesan baru.'
)
ss('[Screenshot 4.10.1: Daftar Chat - Dari Dashboard, tap ikon chat (forum_outlined) di AppBar kanan. Halaman "Direct Message" menampilkan daftar percakapan dengan avatar, nama, pesan terakhir, waktu, dan badge unread.]')
ss('[Screenshot 4.10.2: Detail Chat - Tap salah satu nama di daftar chat. Halaman chat menampilkan bubble pesan (biru = pesan kita, putih = pesan lawan), dengan input bar di bawah berisi text field dan tombol send bulat gradient.]')

# 4.11 Profil Pengguna Lain
doc.add_heading('4.11 Profil Pengguna Lain (User Profile)', level=2)
doc.add_paragraph(
    'Halaman profil pengguna lain memiliki tampilan serupa Profil Saya, '
    'namun dengan tambahan tombol "Ikuti"/"Mengikuti" dan tombol "Pesan" untuk mengirim DM. '
    'Tombol report (flag) tersedia di AppBar untuk melaporkan pengguna. '
    'Jika pengguna adalah admin/superadmin, akan tampil verified badge (centang biru) di samping nama.'
)
ss('[Screenshot 4.11: Profil Pengguna Lain - Dari halaman Forum, tap username di salah satu card diskusi. Halaman profil pengguna lain muncul dengan tombol "Ikuti" dan "Pesan" di bawah nama, serta stats Level & Poin.]')

# 4.12 Leaderboard
doc.add_heading('4.12 Leaderboard (Papan Peringkat)', level=2)
doc.add_paragraph(
    'Halaman Leaderboard menampilkan peringkat kontributor terbaik. '
    'AppBar gradient orange dengan emoji trophy. '
    'Top 3 ditampilkan dengan card spesial: medal emoji dan warna gold (peringkat 1), silver (peringkat 2), bronze (peringkat 3). '
    'Peringkat 4 dan seterusnya ditampilkan dalam numbered list dengan avatar, username, level, dan total points.'
)
ss('[Screenshot 4.12: Leaderboard - Dari Dashboard, scroll ke card "Papan Peringkat" berwarna orange dan tap. Halaman Top Contributors menampilkan 3 teratas dengan card spesial bergradient emas/perak/perunggu, diikuti list ranking berikutnya.]')

# 4.13 Dark Mode
doc.add_heading('4.13 Dark Mode', level=2)
doc.add_paragraph(
    'Seluruh aplikasi Lognity mendukung dark mode yang dapat diaktifkan melalui toggle '
    'ikon bulan/matahari di AppBar Dashboard. Saat dark mode aktif, warna berubah: '
    'background menjadi gelap (slate), card menggunakan warna gelap, dan teks menjadi terang. '
    'Preferensi tema disimpan menggunakan SharedPreferences sehingga persisten antar sesi.'
)
ss('[Screenshot 4.13: Dark Mode - Di Dashboard, tap ikon bulan (dark_mode) di AppBar bagian kanan (sebelum ikon chat). Seluruh tampilan berubah ke mode gelap. Screenshot Dashboard dalam dark mode untuk menunjukkan perbandingan.]')

# 4.14 CRUD Operations
doc.add_heading('4.14 CRUD Operations Summary', level=2)
doc.add_paragraph('Berikut adalah ringkasan operasi CRUD yang tersedia dalam aplikasi:')

table2 = doc.add_table(rows=5, cols=2)
table2.style = 'Table Grid'
h = table2.rows[0].cells
h[0].text = 'Operasi'
h[1].text = 'Detail'
for cell in h:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

crud = [
    ('CREATE', 'Buat forum baru, Registrasi user, Kirim pesan chat, Tambah komentar/jawaban'),
    ('READ', 'Lihat daftar forum, Detail forum, E-Library, Dashboard stats, Leaderboard, Profil user, Chat messages'),
    ('UPDATE', 'Edit forum, Edit komentar, Edit profil (username, email, foto), Ganti password'),
    ('DELETE', 'Hapus forum, Hapus komentar, Hapus chat'),
]
for i, (op, detail) in enumerate(crud):
    row = table2.rows[i + 1].cells
    row[0].text = op
    row[1].text = detail

doc.add_page_break()

# BAB 5 - LAMPIRAN
doc.add_heading('BAB 5 - LAMPIRAN', level=1)
doc.add_paragraph('Link Projek GitHub: https://github.com/aflahzaki/Aplikasi-Perangkat-Bergerak---Lognity')
doc.add_paragraph('Link APK: [Akan dilengkapi]')
doc.add_paragraph('Link Showcase: [Akan dilengkapi]')

# SAVE
output_path = '/projects/sandbox/Aplikasi-Perangkat-Bergerak---Lognity/Laporan_UAS_Lognity.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
